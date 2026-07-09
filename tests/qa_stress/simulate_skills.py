import os
import random
import string
import time

import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor
from fpdf import FPDF
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

OUTPUT_DIR = "test-results/skills_test"
os.makedirs(OUTPUT_DIR, exist_ok=True)

def generate_random_string(length=10):
    """Generates a random string with mixed characters."""
    chars = string.ascii_letters + string.digits + "!@#$%^&*()_+"
    return ''.join(random.choice(chars) for _ in range(length))

def test_chart_skill():
    print("Testing Chart Skill (Advanced)...")
    try:
        # 1. Sine Wave
        x = np.linspace(0, 10, 100)
        y = np.sin(x)
        plt.figure(figsize=(10, 6))
        plt.plot(x, y, label='Sine Wave', color='blue', linewidth=2)
        plt.title('Adversarial Test Chart - Sine Wave')
        plt.xlabel('X Axis')
        plt.ylabel('Y Axis')
        plt.legend()
        plt.grid(True)
        chart1_path = os.path.join(OUTPUT_DIR, "test_chart_sine.png")
        plt.savefig(chart1_path)
        plt.close()

        # 2. Scatter Plot with random data
        x_scatter = np.random.rand(100)
        y_scatter = np.random.rand(100)
        colors = np.random.rand(100)
        sizes = 1000 * np.random.rand(100)
        plt.figure(figsize=(10, 6))
        plt.scatter(x_scatter, y_scatter, c=colors, s=sizes, alpha=0.5, cmap='viridis')
        plt.title('Adversarial Test Chart - Scatter')
        plt.colorbar()
        chart2_path = os.path.join(OUTPUT_DIR, "test_chart_scatter.png")
        plt.savefig(chart2_path)
        plt.close()

        print(f"✅ Charts created: {chart1_path}, {chart2_path}")
        return [chart1_path, chart2_path]
    except Exception as e:
        print(f"❌ Chart creation failed: {e}")
        return []

def test_docx_skill(image_paths):
    print("Testing DOCX Skill (Advanced)...")
    try:
        doc = Document()

        # Title
        title = doc.add_heading('Jules QA Protocol - Advanced DOCX Test', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Introduction
        doc.add_paragraph('This document tests the limits of the python-docx library.')

        # Section 1: Text Styles
        doc.add_heading('1. Text Styles & Formatting', level=1)
        p = doc.add_paragraph()
        p.add_run('Bold text. ').bold = True
        p.add_run('Italic text. ').italic = True
        p.add_run('Underlined text. ').underline = True
        run = p.add_run('Colored text.')
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        # Section 2: Large Table (Stress Test)
        doc.add_heading('2. Large Table (100 Rows)', level=1)
        table = doc.add_table(rows=101, cols=4)
        table.style = 'Table Grid'

        # Header
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'ID'
        hdr_cells[1].text = 'Random String'
        hdr_cells[2].text = 'Value'
        hdr_cells[3].text = 'Status'

        # Data
        for i in range(1, 101):
            row_cells = table.rows[i].cells
            row_cells[0].text = str(i)
            row_cells[1].text = generate_random_string(15)
            row_cells[2].text = str(random.randint(1, 1000))
            row_cells[3].text = random.choice(['Active', 'Inactive', 'Pending', 'Error'])

        # Section 3: Images
        doc.add_heading('3. Embedded Images', level=1)
        for img_path in image_paths:
            if os.path.exists(img_path):
                doc.add_paragraph(f"Image: {os.path.basename(img_path)}")
                doc.add_picture(img_path, width=Inches(5))

        # Section 4: Special Characters
        doc.add_heading('4. Special Characters & Unicode', level=1)
        doc.add_paragraph('Emoji: 🚀 ✅ ❌ 🐍 🐳')
        doc.add_paragraph('Languages: English, Español, Français, Deutsch, 日本語, 中文')
        doc.add_paragraph('Symbols: © ® ™ § ¶ † ‡')

        doc_path = os.path.join(OUTPUT_DIR, "test_doc_advanced.docx")
        doc.save(doc_path)
        print(f"✅ Advanced DOCX created at {doc_path}")
        return doc_path
    except Exception as e:
        print(f"❌ DOCX creation failed: {e}")
        return None

def test_xlsx_skill():
    print("Testing XLSX Skill (Advanced)...")
    try:
        wb = Workbook()

        # Sheet 1: Large Data
        ws1 = wb.active
        ws1.title = "Large Dataset"

        headers = ["ID", "Name", "Category", "Q1", "Q2", "Q3", "Q4", "Total", "Average"]
        ws1.append(headers)

        # Style header
        for cell in ws1[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill(start_color="0000FF", end_color="0000FF", fill_type="solid")
            cell.alignment = Alignment(horizontal="center")

        # Generate 10,000 rows
        print("  Generating 10,000 rows in Excel...")
        for i in range(1, 10001):
            q1 = random.randint(0, 100)
            q2 = random.randint(0, 100)
            q3 = random.randint(0, 100)
            q4 = random.randint(0, 100)
            row = [
                i,
                f"Item-{i}",
                random.choice(["A", "B", "C", "D"]),
                q1, q2, q3, q4,
                f"=SUM(D{i+1}:G{i+1})",
                f"=AVERAGE(D{i+1}:G{i+1})"
            ]
            ws1.append(row)

        # Sheet 2: Formulas & validation
        ws2 = wb.create_sheet(title="Summary")
        ws2['A1'] = "Metric"
        ws2['B1'] = "Value"
        ws2['A2'] = "Total Items"
        ws2['B2'] = "=COUNT('Large Dataset'!A:A)"
        ws2['A3'] = "Average Q1"
        ws2['B3'] = "=AVERAGE('Large Dataset'!D:D)"

        xlsx_path = os.path.join(OUTPUT_DIR, "test_sheet_advanced.xlsx")
        wb.save(xlsx_path)
        print(f"✅ Advanced XLSX created at {xlsx_path}")
        return xlsx_path
    except Exception as e:
        print(f"❌ XLSX creation failed: {e}")
        return None

def test_pdf_skill():
    print("Testing PDF Skill (Advanced)...")
    try:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        # Header
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(0, 10, txt="Jules QA Protocol - Advanced PDF Test", ln=1, align="C")
        pdf.ln(10)

        # Content Loop
        pdf.set_font("Arial", size=10)
        for i in range(1, 51):
            pdf.cell(0, 10, txt=f"Line {i}: This creates a multi-page document to stress test pagination. " + generate_random_string(20), ln=1)

        # Graphics
        pdf.add_page()
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, txt="Graphics Page", ln=1)

        # Draw some lines and rects
        pdf.set_line_width(1)
        pdf.set_draw_color(255, 0, 0)
        pdf.line(10, 40, 200, 40)

        pdf.set_fill_color(0, 0, 255)
        pdf.rect(10, 50, 50, 50, 'FD')

        pdf_path = os.path.join(OUTPUT_DIR, "test_doc_advanced.pdf")
        pdf.output(pdf_path)
        print(f"✅ Advanced PDF created at {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"❌ PDF creation failed: {e}")
        return None

def test_frontend_design_skill():
    print("Testing Frontend Design Skill (Artifact Simulation)...")
    try:
        # Simulate React Component Generation
        react_code = """
import React, { useState } from 'react';
import { Card, CardContent } from '@/components/ui/card';
import { Button } from '@/components/ui/button';
import { LineChart, Line, XAxis, YAxis, Tooltip } from 'recharts';

export default function Dashboard() {
  const [count, setCount] = useState(0);
  const data = [
    {name: 'A', uv: 400},
    {name: 'B', uv: 300},
    {name: 'C', uv: 200}
  ];

  return (
    <div className="p-4 space-y-4">
      <h1 className="text-2xl font-bold">Analytics Dashboard</h1>
      <div className="grid grid-cols-1 md:grid-cols-2 gap-4">
        <Card>
          <CardContent>
            <h2 className="text-xl">User Counter</h2>
            <p className="text-4xl font-bold">{count}</p>
            <Button onClick={() => setCount(count + 1)}>Increment</Button>
          </CardContent>
        </Card>
        <Card>
          <CardContent>
            <LineChart width={300} height={200} data={data}>
              <XAxis dataKey="name" />
              <YAxis />
              <Tooltip />
              <Line type="monotone" dataKey="uv" stroke="#8884d8" />
            </LineChart>
          </CardContent>
        </Card>
      </div>
    </div>
  );
}
"""
        # Validate syntax (basic check)
        if "import React" not in react_code or "export default" not in react_code:
            raise ValueError("Invalid React code structure")

        output_path = os.path.join(OUTPUT_DIR, "test_artifact.jsx")
        with open(output_path, "w") as f:
            f.write(react_code)

        print(f"✅ Frontend Artifact simulated at {output_path}")
        return output_path
    except Exception as e:
        print(f"❌ Frontend Design failed: {e}")
        return None

if __name__ == "__main__":
    print("🚀 Starting Advanced Skills Simulation...")
    start_time = time.time()

    chart_paths = test_chart_skill()
    test_docx_skill(chart_paths)
    test_xlsx_skill()
    test_pdf_skill()
    test_frontend_design_skill()

    duration = time.time() - start_time
    print(f"🏁 Skills Simulation Complete in {duration:.2f} seconds.")
